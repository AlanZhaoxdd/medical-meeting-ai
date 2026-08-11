from __future__ import annotations

import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from app.schemas.export import TextExportSection, TextPreviewRead
from app.services.export_bundle import (
    AnalysisBundle,
    source_index,
)


SECTION_HEADERS: list[tuple[str, str]] = [
    ("overview", "会议概述"),
    ("divergence", "分歧与焦虑"),
    ("evidence", "循证数据解读"),
    ("clinical", "临床用药建议"),
    ("consensus", "专家共识"),
    ("actions", "行动计划"),
]

SECTION_ALIASES: dict[str, str] = {
    "会议概述": "overview",
    "会议总述": "summary",
    "会议概况": "overview",
    "分歧与焦虑": "divergence",
    "分歧与遗留问题": "divergence",
    "分歧": "divergence",
    "循证数据解读": "evidence",
    "临床用药建议": "clinical",
    "专家共识": "consensus",
    "核心结论与共识": "consensus",
    "共识": "consensus",
    "行动计划": "actions",
    "行动项": "actions",
}


def _known_header(title: str) -> str | None:
    cleaned = re.sub(r"^#+\s*", "", title).strip()
    cleaned = re.sub(r"^\*\*(.+?)\*\*$", r"\1", cleaned).strip()
    cleaned = re.sub(r"^[一二三四五六七八九十]+、\s*", "", cleaned)
    for alias, key in SECTION_ALIASES.items():
        if cleaned == alias or cleaned.startswith(alias):
            return key
    return None


def split_minutes_sections(content: str) -> list[tuple[str, str, list[int]]]:
    """Split the confirmed AI minutes markdown into (key, body, citations)."""

    sections: list[tuple[str, str, list[int]]] = []
    if not content:
        return sections
    heading_pattern = re.compile(
        r"(?m)^(?:#{2,6}\s+(.+?)\s*|\*\*([一二三四五六七八九十]+、[^*\n]+?)\*\*\s*)$"
    )
    matches = list(heading_pattern.finditer(content))
    pending_intro = content[: matches[0].start()].strip() if matches else content.strip()
    if pending_intro:
        sections.append(("intro", pending_intro, []))
    for index, match in enumerate(matches):
        title = (match.group(1) or match.group(2) or "").strip()
        body_start = match.end()
        body_end = matches[index + 1].start() if index + 1 < len(matches) else len(content)
        body = content[body_start:body_end].strip()
        citations = [
            int(value)
            for value in re.findall(r"\[(\d+)\]", body)
            if value.isdigit()
        ]
        sections.append((title, body, citations))
    return sections


def _strip_markdown(value: str) -> str:
    text = value
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "• ", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def _escape_paragraph(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _line_items(body: str) -> list[str]:
    lines: list[str] = []
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^[-*+•]\s+", stripped):
            lines.append(_strip_markdown(re.sub(r"^[-*+•]\s+", "", stripped)))
        elif re.match(r"^\d+[.、]\s+", stripped):
            lines.append(_strip_markdown(re.sub(r"^\d+[.、]\s+", "", stripped)))
    return lines


def _remove_citation_markers(value: str) -> str:
    return re.sub(r"\s*\[\d+\]", "", value).strip()


def _parse_minutes(bundle: AnalysisBundle) -> dict[str, tuple[str, list[int]]]:
    """Map confirmed minutes markdown to canonical section keys."""

    minutes_module = next(
        (module for module in bundle.modules if module.get("id") == "minutes"),
        None,
    )
    content = str((minutes_module or {}).get("content") or "")
    parsed: list[tuple[str, str, list[int]]] = []
    for title, body, citations in split_minutes_sections(content):
        key = _known_header(title)
        parsed.append((key or title, body, citations))
    result: dict[str, tuple[str, list[int]]] = {}
    for key, body, citations in parsed:
        if key in SECTION_HEADERS_MAP and key not in result:
            result[key] = (body, citations)
    return result


SECTION_HEADERS_MAP = {key: title for key, title in SECTION_HEADERS}


def compose_export_sections(
    bundle: AnalysisBundle,
    *,
    show_attendee_names: bool,
    include_references: bool,
    include_citation_markers: bool,
) -> list[TextExportSection]:
    """Build the ordered export sections from the confirmed analysis only."""

    minutes = _parse_minutes(bundle)
    sources = bundle.sources

    def make(
        key: str,
        title: str,
        body: str | None,
        citations: list[int],
        items: list[str] | None = None,
    ) -> TextExportSection:
        cleaned = body.strip() if body else ""
        if items is None:
            items = _line_items(cleaned)
        return TextExportSection(
            key=key,
            title=title,
            content=cleaned if not items else cleaned,
            items=items,
            citations=sorted(set(c for c in citations if 1 <= c <= len(sources))),
        )

    available: dict[str, TextExportSection] = {}
    for key, title in SECTION_HEADERS:
        if body := minutes.get(key):
            available[key] = make(key, title, body[0], body[1])

    ordered_keys = [key for key, _ in SECTION_HEADERS]
    sections: list[TextExportSection] = []
    for key in ordered_keys:
        section = available.get(key)
        if section is None:
            continue
        sections.append(section)

    if include_references and sources:
        sections.append(
            TextExportSection(
                key="sources",
                title="引用来源或知识库依据",
                content=None,
                items=[
                    f"[{source_index(item)}] {item.get('title', '来源')}："
                    f"{str(item.get('snippet') or '')[:200]}"
                    for item in sources
                ],
                citations=[],
            )
        )

    if not include_citation_markers:
        for section in sections:
            if section.content:
                section.content = _remove_citation_markers(section.content)
            section.items = [_remove_citation_markers(item) for item in section.items]

    return sections


def default_file_name(bundle: AnalysisBundle, fmt: str) -> str:
    safe = re.sub(r"[\\/:*?\"<>|\s]+", "-", bundle.meeting.title).strip("-")
    safe = (safe or "会议成果")[:60]
    return f"{safe}-AI会议纪要-{datetime.now(timezone.utc).strftime('%Y%m%d')}.{fmt}"


def build_text_preview(
    bundle: AnalysisBundle,
    *,
    show_attendee_names: bool,
    include_cover: bool,
    include_references: bool,
    include_citation_markers: bool,
) -> TextPreviewRead:
    sections = compose_export_sections(
        bundle,
        show_attendee_names=show_attendee_names,
        include_references=include_references,
        include_citation_markers=include_citation_markers,
    )
    return TextPreviewRead(
        meeting_id=bundle.meeting.id,
        meeting_title=bundle.meeting.title,
        starts_at=bundle.meeting.starts_at.isoformat() if bundle.meeting.starts_at else None,
        ends_at=bundle.meeting.ends_at.isoformat() if bundle.meeting.ends_at else None,
        location=bundle.meeting.location,
        organizer=bundle.meeting.organizer,
        topic=bundle.meeting.topic,
        analysis_version=bundle.analysis_version,
        include_cover=include_cover,
        sections=sections,
        sources=[
            {
                "index": source_index(item),
                "title": item.get("title"),
                "snippet": item.get("snippet"),
                "type": item.get("type"),
            }
            for item in bundle.sources
        ],
    )


def _docx_body_text(sections: list[TextExportSection]) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for section in sections:
        if section.items:
            blocks.append((section.title, "\n".join(f"• {item}" for item in section.items)))
        elif section.content:
            blocks.append((section.title, _strip_markdown(section.content)))
    return blocks


def render_text_docx(
    bundle: AnalysisBundle,
    *,
    include_cover: bool,
    sections: list[TextExportSection],
    include_references: bool,
) -> bytes:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Pt

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = section.page_width
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Microsoft YaHei"
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(4)
    heading_style.paragraph_format.keep_with_next = True

    if include_cover:
        document.add_heading(bundle.meeting.title, level=0)
        document.add_paragraph("会议成果导出 · 文字版会议纪要")
        if bundle.meeting.starts_at:
            document.add_paragraph(
                f"会议日期：{bundle.meeting.starts_at.strftime('%Y-%m-%d %H:%M')}"
            )
        if bundle.meeting.organizer:
            document.add_paragraph(f"组织方：{bundle.meeting.organizer}")
        if bundle.meeting.location:
            document.add_paragraph(f"地点：{bundle.meeting.location}")

    body_sections = sections
    if not include_references:
        body_sections = [section for section in sections if section.key != "sources"]
    for title, body in _docx_body_text(body_sections):
        document.add_heading(title, level=1)
        for paragraph_text in body.splitlines():
            document.add_paragraph(paragraph_text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_text_pdf(
    bundle: AnalysisBundle,
    *,
    include_cover: bool,
    sections: list[TextExportSection],
    include_references: bool,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    base = ParagraphStyle(
        "base",
        fontName="STSong-Light",
        fontSize=10.5,
        leading=17,
        wordWrap="CJK",
    )
    heading = ParagraphStyle(
        "heading",
        parent=base,
        fontSize=14,
        leading=20,
        spaceBefore=12,
        spaceAfter=6,
        textColor="#123c53",
    )
    title_style = ParagraphStyle(
        "title",
        parent=base,
        fontSize=22,
        leading=30,
        spaceAfter=18,
        alignment=1,
        textColor="#123c53",
    )
    meta = ParagraphStyle("meta", parent=base, fontSize=9.5, leading=14)

    story: list[Any] = []
    if include_cover:
        story.append(Paragraph(bundle.meeting.title, title_style))
        story.append(Paragraph("会议成果导出 · 文字版会议纪要", base))
        story.append(Spacer(1, 10 * mm))
        if bundle.meeting.starts_at:
            story.append(
                Paragraph(
                    f"会议日期：{bundle.meeting.starts_at.strftime('%Y-%m-%d %H:%M')}",
                    meta,
                )
            )
        if bundle.meeting.organizer:
            story.append(Paragraph(f"组织方：{bundle.meeting.organizer}", meta))
        if bundle.meeting.location:
            story.append(Paragraph(f"地点：{bundle.meeting.location}", meta))
        story.append(Spacer(1, 14 * mm))

    body_sections = sections
    if not include_references:
        body_sections = [section for section in sections if section.key != "sources"]
    for section in body_sections:
        story.append(Paragraph(section.title, heading))
        if section.items:
            for item in section.items:
                story.append(Paragraph(f"• {_escape_paragraph(item)}", base))
        elif section.content:
            for paragraph in _strip_markdown(section.content).splitlines():
                if paragraph.strip():
                    story.append(Paragraph(_escape_paragraph(paragraph), base))
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    document.build(story)
    return buffer.getvalue()


def render_text_file(
    bundle: AnalysisBundle,
    *,
    fmt: str,
    include_cover: bool,
    show_attendee_names: bool,
    include_references: bool,
    include_citation_markers: bool,
) -> tuple[bytes, str]:
    sections = compose_export_sections(
        bundle,
        show_attendee_names=show_attendee_names,
        include_references=include_references,
        include_citation_markers=include_citation_markers,
    )
    if fmt == "docx":
        content = render_text_docx(
            bundle,
            include_cover=include_cover,
            sections=sections,
            include_references=include_references,
        )
        return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    content = render_text_pdf(
        bundle,
        include_cover=include_cover,
        sections=sections,
        include_references=include_references,
    )
    return content, "application/pdf"
